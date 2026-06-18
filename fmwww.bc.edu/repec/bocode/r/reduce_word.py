#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
reduce_word.py - AIGC Similarity Reduction Tool (Bilingual)
Version: 8.1 - Stable (No NLTK auto-loading)
Date: 06June2026
"""

import argparse
import random
import re
import json
import os
import sys
import io
from docx import Document

# ============================================================================
# Fix Windows console encoding
# ============================================================================

if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')
else:
    if sys.stdout.encoding != 'utf-8':
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    if sys.stderr.encoding != 'utf-8':
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

# ============================================================================
# Import Chinese NLP (optional)
# ============================================================================

try:
    import jieba
    JIEBA_AVAILABLE = True
except ImportError:
    JIEBA_AVAILABLE = False
    print(u"提示: jieba未安装，中文模式使用基础方法。安装: pip install jieba")

# ============================================================================
# English synonyms - Built-in dictionary (no NLTK required)
# ============================================================================

EN_SYNONYMS = {
    "important": ["significant", "crucial", "vital", "essential", "critical", "paramount"],
    "research": ["study", "investigation", "analysis", "examination", "exploration"],
    "method": ["approach", "technique", "procedure", "methodology", "process"],
    "result": ["outcome", "finding", "consequence", "product", "achievement"],
    "show": ["demonstrate", "indicate", "reveal", "exhibit", "display", "illustrate"],
    "increase": ["rise", "grow", "elevate", "raise", "boost", "escalate"],
    "decrease": ["reduce", "decline", "diminish", "lower", "drop", "fall"],
    "good": ["positive", "favorable", "beneficial", "advantageous", "desirable"],
    "bad": ["negative", "adverse", "detrimental", "harmful", "unfavorable"],
    "big": ["large", "substantial", "considerable", "significant", "massive"],
    "small": ["tiny", "minor", "modest", "limited", "slight"],
    "many": ["numerous", "countless", "multiple", "various", "abundant"],
    "develop": ["evolve", "advance", "progress", "formulate", "establish"],
    "analyze": ["examine", "assess", "evaluate", "study", "investigate"],
    "provide": ["supply", "offer", "furnish", "deliver", "present"],
    "need": ["require", "demand", "necessitate", "call for"],
    "use": ["utilize", "employ", "apply", "leverage", "adopt"],
    "make": ["create", "produce", "generate", "construct", "form"],
    "get": ["obtain", "acquire", "gain", "secure", "receive"],
    "change": ["alter", "modify", "adjust", "transform", "shift"],
    "significant": ["substantial", "considerable", "notable", "remarkable", "striking"],
    "different": ["various", "diverse", "distinct", "dissimilar", "contrasting"],
    "similar": ["alike", "comparable", "analogous", "corresponding", "equivalent"],
    "cause": ["lead to", "result in", "bring about", "give rise to", "produce"],
    "effect": ["impact", "influence", "consequence", "outcome", "ramification"],
    "improve": ["enhance", "boost", "strengthen", "optimize", "refine"],
    "reduce": ["decrease", "lower", "diminish", "cut", "minimize"],
    "analysis": ["examination", "assessment", "evaluation", "investigation", "review"],
    "data": ["information", "statistics", "figures", "material", "evidence"],
    "process": ["procedure", "method", "approach", "technique", "operation"],
    "system": ["framework", "structure", "network", "organization", "setup"],
    "model": ["framework", "representation", "simulation", "prototype", "design"],
    "theory": ["concept", "principle", "hypothesis", "framework", "explanation"],
}

# Chinese synonyms
ZH_SYNONYMS = {
    "研究": ["探讨", "考察", "分析", "探究", "钻研"],
    "方法": ["方式", "手段", "途径", "策略", "方案"],
    "重要": ["关键", "核心", "主要", "首要", "重大"],
    "结果": ["成果", "结论", "效果", "成效", "产物"],
    "显示": ["表明", "说明", "揭示", "展现", "呈现"],
    "增加": ["增长", "上升", "提高", "增强", "提升"],
    "减少": ["降低", "下降", "削减", "缩减", "减轻"],
    "好": ["良好", "优秀", "出色", "优越", "优良"],
    "坏": ["不良", "负面", "有害", "不利", "糟糕"],
    "大": ["巨大", "庞大", "宏大", "广阔", "重大"],
    "小": ["微小", "细小", "少量", "有限", "轻微"],
    "许多": ["众多", "大量", "无数", "诸多", "繁多"],
    "发展": ["进步", "演进", "推进", "壮大", "成长"],
    "分析": ["剖析", "解析", "研判", "评估", "考核"],
    "提供": ["给予", "供给", "供应", "馈赠", "献出"],
    "需要": ["需求", "要求", "须要", "必要"],
    "使用": ["运用", "采用", "利用", "应用", "借重"],
    "制作": ["创作", "生产", "生成", "构建", "制造"],
    "获得": ["取得", "赢得", "争取", "获取", "博得"],
    "改变": ["转变", "转换", "调整", "变革", "改动"],
}

# ============================================================================
# Text cleaning functions
# ============================================================================

def clean_text_for_xml(text):
    if text is None:
        return u""
    text = text.replace(u'\x00', u'')
    
    def is_valid_xml_char(c):
        code = ord(c)
        return (code == 0x09 or code == 0x0A or code == 0x0D or
                (0x20 <= code <= 0xD7FF) or (0xE000 <= code <= 0xFFFD) or
                (0x10000 <= code <= 0x10FFFF))
    
    return u''.join(c for c in text if is_valid_xml_char(c))

def safe_set_paragraph_text(paragraph, text):
    try:
        cleaned_text = clean_text_for_xml(text)
        paragraph.clear()
        if cleaned_text:
            paragraph.add_run(cleaned_text)
        return True
    except Exception as e:
        try:
            import string
            printable = set(string.printable)
            safe_text = u''.join(c for c in text if c in printable)
            paragraph.clear()
            if safe_text:
                paragraph.add_run(safe_text)
            return True
        except:
            return False

# ============================================================================
# Reduction class
# ============================================================================

class TextReducer:
    def __init__(self, language='zh', intensity=0.5, user_dict=None, methods=None):
        self.language = language
        self.intensity = intensity
        self.user_dict = user_dict or {}
        self.methods = methods or ['synonym', 'voice', 'neg', 'split', 'order', 'expand', 'modifier']
        
        self.syn_dict = {}
        if language == 'zh':
            self.syn_dict.update(ZH_SYNONYMS)
        else:
            self.syn_dict.update(EN_SYNONYMS)
        
        for key, values in self.user_dict.items():
            self.syn_dict[key] = values
    
    def synonym_replace(self, text):
        if not text:
            return text, 0
        
        words_replaced = 0
        result = text
        
        if self.language == 'zh':
            for word, synonyms in self.syn_dict.items():
                if word in result and random.random() < self.intensity:
                    synonym = random.choice(synonyms)
                    result = result.replace(word, synonym)
                    words_replaced += 1
        else:
            # English: word boundary replacement
            words = re.findall(r'\b\w+\b|[^\w\s]|\s+', text)
            new_words = []
            
            for token in words:
                if re.match(r'^\w+$', token):
                    token_lower = token.lower()
                    if token_lower in self.syn_dict and random.random() < self.intensity:
                        synonym = random.choice(self.syn_dict[token_lower])
                        if token[0].isupper():
                            synonym = synonym.capitalize()
                        elif token.isupper():
                            synonym = synonym.upper()
                        new_words.append(synonym)
                        words_replaced += 1
                    else:
                        new_words.append(token)
                else:
                    new_words.append(token)
            
            result = ''.join(new_words)
        
        return result, words_replaced
    
    def voice_change(self, text):
        if not text or random.random() > self.intensity * 1.2:
            return text, 0
        
        changes = 0
        result = text
        
        if self.language == 'zh':
            patterns = [(u'(\w+)执行了(\w+)', u'\2被\1执行'),
                       (u'(\w+)完成了(\w+)', u'\2由\1完成')]
            for pattern, replacement in patterns:
                try:
                    if re.search(pattern, result):
                        result = re.sub(pattern, replacement, result)
                        changes += 1
                        break
                except:
                    pass
        else:
            patterns = [(r'(\w+) (performs|conducts) ([\w\s]+)', r'\3 is \2ed by \1'),
                       (r'(\w+) (shows|indicates) ([\w\s]+)', r'\3 is \2d by \1')]
            for pattern, replacement in patterns:
                try:
                    if re.search(pattern, result, re.I):
                        result = re.sub(pattern, replacement, result, flags=re.I)
                        changes += 1
                        break
                except:
                    pass
        
        return result, changes
    
    def neg_inversion(self, text):
        if not text or random.random() > self.intensity * 0.8:
            return text, 0
        
        changes = 0
        result = text
        
        if self.language == 'zh':
            patterns = [(u'很重要', u'不是不重要'), (u'非常必要', u'并非不必要')]
            for orig, repl in patterns:
                if orig in result:
                    result = result.replace(orig, repl)
                    changes += 1
                    break
        else:
            patterns = [(r'important', 'not unimportant'), (r'necessary', 'not unnecessary')]
            for orig, repl in patterns:
                try:
                    if re.search(r'\b' + orig + r'\b', result, re.I):
                        result = re.sub(r'\b' + orig + r'\b', repl, result, flags=re.I)
                        changes += 1
                        break
                except:
                    pass
        
        return result, changes
    
    def split_merge(self, text):
        if not text or random.random() > self.intensity:
            return text, 0
        
        changes = 0
        result = text
        
        if self.language == 'zh':
            if len(text) > 30 and u'，' in text:
                parts = text.split(u'，', 1)
                if len(parts) == 2:
                    result = parts[0] + u'。' + parts[1]
                    changes += 1
        else:
            if len(text) > 80 and ', ' in text:
                parts = text.split(', ', 1)
                if len(parts) == 2:
                    result = parts[0] + '. ' + parts[1]
                    changes += 1
            elif len(text) < 40 and '. ' in text:
                result = text.replace('. ', ', ', 1)
                changes += 1
        
        return result, changes
    
    def order_adjust(self, text):
        if not text or random.random() > self.intensity * 0.9:
            return text, 0
        
        changes = 0
        result = text
        
        if self.language == 'zh':
            patterns = [(u'(因为.*?)(所以.*?)', u'\2，\1')]
            for pattern, replacement in patterns:
                try:
                    if re.search(pattern, result):
                        result = re.sub(pattern, replacement, result)
                        changes += 1
                        break
                except:
                    pass
        else:
            patterns = [(r'(Because.*?),(.*?)', r'\2 because \1')]
            for pattern, replacement in patterns:
                try:
                    if re.search(pattern, result, re.I):
                        result = re.sub(pattern, replacement, result, flags=re.I)
                        changes += 1
                        break
                except:
                    pass
        
        return result, changes
    
    def expand_abbr(self, text):
        if not text or self.language != 'en' or random.random() > self.intensity * 0.7:
            return text, 0
        
        changes = 0
        result = text
        abbr_dict = {"e.g.": "for example", "i.e.": "that is", "etc.": "and so on",
                    "vs.": "versus", "Mr.": "Mister", "Dr.": "Doctor"}
        
        for abbr, full in abbr_dict.items():
            if abbr in result:
                result = result.replace(abbr, full)
                changes += 1
        
        return result, changes
    
    def add_modifier(self, text):
        if not text or random.random() > self.intensity * 0.6:
            return text, 0
        
        changes = 0
        result = text
        
        if self.language == 'zh':
            modifiers = [u'显著地', u'明显地', u'尤其', u'特别地']
        else:
            modifiers = ['significantly', 'notably', 'particularly', 'importantly']
        
        if len(text) > 20 and random.random() < 0.3:
            modifier = random.choice(modifiers)
            result = (modifier + u'，' + result) if self.language == 'zh' else (modifier + ', ' + result)
            changes += 1
        
        return result, changes
    
    def reduce_text(self, text):
        if not text or len(text.strip()) < 5:
            return text, 0
        
        current_text = text
        total_changes = 0
        
        method_functions = {
            'synonym': self.synonym_replace,
            'voice': self.voice_change,
            'neg': self.neg_inversion,
            'split': self.split_merge,
            'order': self.order_adjust,
            'expand': self.expand_abbr,
            'modifier': self.add_modifier,
        }
        
        for method in self.methods:
            if method in method_functions:
                try:
                    new_text, changes = method_functions[method](current_text)
                    if changes > 0 and new_text:
                        current_text = new_text
                        total_changes += changes
                except:
                    pass
        
        return current_text, total_changes

# ============================================================================
# Main processing
# ============================================================================

def process_document(input_path, output_path, target=0.3, intensity=0.5, 
                     user_dict=None, methods=None, language='zh'):
    
    print(u"正在加载文档: {}".format(input_path))
    doc = Document(input_path)
    
    reducer = TextReducer(language=language, intensity=intensity,
                         user_dict=user_dict, methods=methods)
    
    print(u"语言模式: {} (使用内置词典)".format('中文' if language=='zh' else '英文'))
    
    total_paragraphs = 0
    modified_paragraphs = 0
    total_changes = 0
    error_paragraphs = 0
    
    # Process paragraphs
    for para in doc.paragraphs:
        if para.text and para.text.strip():
            total_paragraphs += 1
            original_text = para.text
            
            try:
                new_text, changes = reducer.reduce_text(original_text)
                if changes > 0 and new_text and new_text != original_text:
                    if safe_set_paragraph_text(para, new_text):
                        modified_paragraphs += 1
                        total_changes += changes
                    else:
                        error_paragraphs += 1
            except Exception as e:
                error_paragraphs += 1
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text and para.text.strip():
                        total_paragraphs += 1
                        original_text = para.text
                        try:
                            new_text, changes = reducer.reduce_text(original_text)
                            if changes > 0 and new_text and new_text != original_text:
                                if safe_set_paragraph_text(para, new_text):
                                    modified_paragraphs += 1
                                    total_changes += changes
                                else:
                                    error_paragraphs += 1
                        except:
                            error_paragraphs += 1
    
    # Save
    print(u"正在保存文档: {}".format(output_path))
    doc.save(output_path)
    
    # Statistics
    print(u"\n" + "=" * 60)
    print(u"降重完成！")
    print("=" * 60)
    print(u"{:<20} {}".format("语言模式:", "中文" if language=='zh' else "英文"))
    print(u"{:<20} {}".format("总段落数:", total_paragraphs))
    print(u"{:<20} {}".format("修改段落数:", modified_paragraphs))
    if total_paragraphs > 0:
        print(u"{:<20} {:.1f}%".format("修改比例:", modified_paragraphs/total_paragraphs*100))
    print(u"{:<20} {}".format("总修改次数:", total_changes))
    if error_paragraphs > 0:
        print(u"{:<20} {}".format("错误段落数:", error_paragraphs))
    print(u"{:<20} {}".format("使用方法:", ", ".join(reducer.methods)))
    print(u"{:<20} {}".format("输出文件:", output_path))
    print("=" * 60)
    
    return True

def load_user_dict(dict_path):
    if not dict_path or not os.path.exists(dict_path):
        return {}
    
    user_dict = {}
    try:
        if dict_path.endswith('.json'):
            with open(dict_path, 'r', encoding='utf-8') as f:
                user_dict = json.load(f)
        elif dict_path.endswith('.csv'):
            import csv
            with open(dict_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                for row in reader:
                    if len(row) >= 2:
                        user_dict[row[0]] = row[1:]
        else:
            with open(dict_path, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if '=' in line:
                        key, values = line.split('=', 1)
                        user_dict[key.strip()] = [v.strip() for v in values.split(',')]
        
        print(u"已加载用户词典: {} 个条目".format(len(user_dict)))
    except Exception as e:
        print(u"警告: 加载用户词典失败: {}".format(e))
    
    return user_dict

def main():
    parser = argparse.ArgumentParser(description='Reduce AIGC similarity')
    parser.add_argument('input', help='Input Word document')
    parser.add_argument('-o', '--output', help='Output document')
    parser.add_argument('-t', '--target', type=float, default=0.3)
    parser.add_argument('-i', '--intensity', type=float, default=0.5)
    parser.add_argument('-d', '--dict', dest='dict_path', help='User dictionary')
    parser.add_argument('-m', '--methods', help='Comma-separated methods')
    parser.add_argument('-l', '--language', default='zh', choices=['en', 'zh'])
    
    args = parser.parse_args()
    
    if not os.path.exists(args.input):
        print(u"错误: 文件不存在: {}".format(args.input))
        sys.exit(1)
    
    if args.target < 0.1 or args.target > 0.95:
        print(u"错误: target 必须在 0.1-0.95 之间")
        sys.exit(1)
    
    if args.intensity < 0.1 or args.intensity > 0.95:
        print(u"错误: intensity 必须在 0.1-0.95 之间")
        sys.exit(1)
    
    if not args.output:
        args.output = args.input.replace('.docx', '_reduced.docx') if args.input.endswith('.docx') else args.input + '_reduced.docx'
    
    user_dict = load_user_dict(args.dict_path)
    methods = [m.strip() for m in args.methods.split(',')] if args.methods else None
    
    try:
        process_document(args.input, args.output, args.target, args.intensity,
                        user_dict, methods, args.language)
    except Exception as e:
        print(u"错误: {}".format(e))
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == '__main__':
    main()